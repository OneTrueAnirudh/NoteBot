from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from g4f.client import Client

class TextSummarizer:
    num = ["1. ","2. ","3. ","4. ","5. ","6. ","7. ","8. ","9. ","10. ","11. ","12. ","13. ","14. ","15. "]
    def __init__(self):
        self.client = Client()
    
    def llm_summary(self, keyword, chunks):
        """Queries the LLM to summarize chunks for a specific keyword or a group of keywords, with bullet points."""
        prompt = f"Using the following text segments that describe the keyword {keyword}:\n\n{chunks}\n\nPlease summarize the key points as points. Each point should describe an important takeaway from the provided information. Format your response as a list, but use no other formatting and dont say anything else. Just put the points in the form of a list, where each is an element of a list. Thank you."
        
        while True:
            try:
                response = self.client.chat.completions.create(
                    model="gpt-4o",
                    messages=[{"role": "user", "content": prompt}],
                )
                result = response.choices[0].message.content
                if "Model not found or too long input. Or any other error (xD)" in result:
                    # print("Error: Model not found or input too long, retrying...")
                    continue  
                if "Generated by BLACKBOX.AI, try unlimited chat https://www.blackbox.ai" in result:
                    result = result.replace("Generated by BLACKBOX.AI, try unlimited chat https://www.blackbox.ai", "").strip()
                bullet_points = [point.strip().lstrip("-").strip() for point in result.split("\n") if point.strip()]                
                return bullet_points            
            except Exception as e:
                print(f"An error occurred: {e}. Retrying...")
                continue  
    
    def create_word_document(self, content_dict, name="notes.docx", num=num):
        """Generates a Word document with structured content from LLM, including bullet points."""
        doc = Document()
        for keyword, bullet_points in content_dict.items():
            heading_paragraph = doc.add_paragraph("")
            heading_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            heading_run = heading_paragraph.add_run(keyword)
            heading_run.bold = True
            heading_run.font.name = 'Times New Roman'
            heading_run.font.size = Pt(14)
            for point in bullet_points:
                for i in num:
                    if i in point:
                        point = point[3:]
                paragraph = doc.add_paragraph(point, style='List Bullet')
                paragraph_format = paragraph.style.font
                paragraph_format.size = Pt(13)
                paragraph_format.name = 'Times New Roman'
        doc.save(name)

    def process_all_chunks(self, chunks_by_keyword):
        """Combines all chunks per keyword, generates a summary, and adds to dictionary."""
        summaries = {}
        for keyword, chunks in chunks_by_keyword.items():
            combined_chunks = " ".join(chunks)
            bullet_points = self.llm_summary(keyword, combined_chunks)
            summaries[keyword] = bullet_points
        return summaries
    
    def process_chunks(self, chunks_by_keyword):
        """Processes and summarizes chunks per keyword in a sequential manner."""
        summaries = {}
        for keyword, chunks in chunks_by_keyword.items():
            bullet_points = self.llm_summary(keyword, " ".join(chunks))
            summaries[keyword] = bullet_points
        return summaries


# summarizer = TextSummarizer()

# # Option 1: Process all chunks at once (for each keyword)
# all_chunks_summary = summarizer.process_all_chunks(chunks_by_keyword)
# summarizer.create_word_document(all_chunks_summary, name="all.docx")

# # Option 2: Process each keyword’s chunks separately
# individual_summaries = summarizer.process_chunks(chunks_by_keyword)
# summarizer.create_word_document(individual_summaries, name="ind.docx")
